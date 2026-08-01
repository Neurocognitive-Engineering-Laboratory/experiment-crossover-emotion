from __future__ import annotations

from pathlib import Path
import math

import pandas as pd

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches, Pt

from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import (
    getSampleStyleSheet,
    ParagraphStyle,
)
from reportlab.lib.units import cm
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
    PageBreak,
)


# ======================================================================
# PATHS
# ======================================================================

PROJECT_ROOT = Path(__file__).resolve().parents[1]

OUTPUT_DIR = (
    PROJECT_ROOT
    / "reports"
    / "final_report"
)

OUTPUT_DIR.mkdir(
    parents=True,
    exist_ok=True,
)

DOCX_PATH = (
    OUTPUT_DIR
    / "hypothesis_01_report.docx"
)

PDF_PATH = (
    OUTPUT_DIR
    / "hypothesis_01_report.pdf"
)


# ======================================================================
# STUDY INFORMATION
# ======================================================================

REPORT_TITLE = (
    "Hypothesis 1 - Effect of Emotional Condition "
    "on Working-Memory Performance"
)

HYPOTHESIS = (
    "Positive, negative, and neutral emotional conditions "
    "differentially affect working-memory performance."
)

N_PARTICIPANTS = 23

N_ACCURACY_OBSERVATIONS = 17820

N_RT_OBSERVATIONS = 14184


# ======================================================================
# RAW RESULTS
# ======================================================================

accuracy_descriptive = pd.DataFrame(
    {
        "Condition": [
            "Negative",
            "Neutral",
            "Positive",
        ],
        "Trials": [
            5940,
            5940,
            5940,
        ],
        "Accuracy": [
            0.838047,
            0.850000,
            0.851010,
        ],
        "SD": [
            0.368439,
            0.357101,
            0.356108,
        ],
    }
)


accuracy_model = pd.DataFrame(
    {
        "Contrast": [
            "Negative vs Neutral",
            "Positive vs Neutral",
        ],
        "Beta": [
            -0.0841,
            0.0015,
        ],
        "SE": [
            0.064,
            0.057,
        ],
        "z": [
            -1.313,
            0.027,
        ],
        "p": [
            0.189,
            0.979,
        ],
        "CI_low": [
            -0.210,
            -0.110,
        ],
        "CI_high": [
            0.041,
            0.113,
        ],
    }
)


rt_model = pd.DataFrame(
    {
        "Contrast": [
            "Negative vs Neutral",
            "Positive vs Neutral",
        ],
        "Beta_log_RT": [
            -0.265,
            -0.285,
        ],
        "SE": [
            0.015,
            0.015,
        ],
        "z": [
            -17.288,
            -18.598,
        ],
        "p": [
            0.000001,  # reported as p < .001
            0.000001,
        ],
        "CI_low": [
            -0.295,
            -0.316,
        ],
        "CI_high": [
            -0.235,
            -0.255,
        ],
    }
)


# ======================================================================
# DERIVED EFFECTS
# ======================================================================

accuracy_model["Odds_Ratio"] = (
    accuracy_model["Beta"]
    .apply(math.exp)
)

accuracy_model["OR_CI_low"] = (
    accuracy_model["CI_low"]
    .apply(math.exp)
)

accuracy_model["OR_CI_high"] = (
    accuracy_model["CI_high"]
    .apply(math.exp)
)


rt_model["RT_ratio"] = (
    rt_model["Beta_log_RT"]
    .apply(math.exp)
)

rt_model["Percent_change"] = (
    (
        rt_model["RT_ratio"]
        - 1
    )
    * 100
)

rt_model["Percent_CI_low"] = (
    (
        rt_model["CI_low"]
        .apply(math.exp)
        - 1
    )
    * 100
)

rt_model["Percent_CI_high"] = (
    (
        rt_model["CI_high"]
        .apply(math.exp)
        - 1
    )
    * 100
)


# ======================================================================
# FORMAT HELPERS
# ======================================================================

def format_p(
    p: float,
) -> str:
    """
    Format p-values according to common scientific conventions.
    """

    if p < 0.001:
        return "< .001"

    return f"= {p:.3f}"


def format_ci(
    lower: float,
    upper: float,
    digits: int = 2,
) -> str:

    return (
        f"[{lower:.{digits}f}, "
        f"{upper:.{digits}f}]"
    )


# ======================================================================
# TEXT USED IN BOTH REPORTS
# ======================================================================

METHOD_TEXT = """
A trial-level confirmatory analysis was conducted to evaluate whether
working-memory performance differed across neutral, positive, and negative
emotional conditions. The analytical sample included 23 participants.

Response accuracy was analyzed using a generalized estimating equation
(GEE) model with a binomial distribution, logit link, exchangeable
within-participant correlation structure, and robust standard errors.
Emotional condition was entered as the primary predictor, with the neutral
condition specified as the reference category. The accuracy model included
17,820 trial-level observations clustered within 23 participants.

Reaction time was analyzed among valid correct-response trials only.
Because reaction-time data were positively skewed, reaction time was
log-transformed prior to modeling. A linear mixed-effects model was fitted
with emotional condition as a fixed effect and a participant-specific
random intercept. The reaction-time analysis included 14,184 observations
from 23 participants. Maximum-likelihood estimation was used.

Model coefficients from the binomial GEE were exponentiated and interpreted
as odds ratios. Coefficients from the log reaction-time model were
exponentiated and converted to percentage changes relative to the neutral
condition. Statistical significance was evaluated using a two-sided alpha
level of .05, and 95% confidence intervals are reported.
""".strip()


ACCURACY_RESULTS_TEXT = """
Descriptively, accuracy was highly similar across emotional conditions.
Mean trial-level accuracy was 83.8% in the negative condition, 85.0% in
the neutral condition, and 85.1% in the positive condition.

The omnibus Wald test indicated that emotional condition was not
significantly associated with response accuracy,
chi-square(2) = 2.54, p = .281.

Relative to the neutral condition, the negative condition was associated
with a small, non-significant reduction in the odds of a correct response
(beta = -0.084, SE = 0.064, z = -1.31, p = .189), corresponding to an
odds ratio of approximately 0.92 (95% CI [0.81, 1.04]).

The positive condition did not differ from the neutral condition
(beta = 0.002, SE = 0.057, z = 0.03, p = .979), with an odds ratio of
approximately 1.00 (95% CI [0.90, 1.12]).

Taken together, these findings provide no evidence that emotional
condition substantially altered the probability of producing a correct
response.
""".strip()


RT_RESULTS_TEXT = """
In contrast to accuracy, emotional condition was strongly associated with
reaction time.

Relative to the neutral condition, the negative emotional condition was
associated with a significantly lower log reaction time
(beta = -0.265, SE = 0.015, z = -17.29, p < .001,
95% CI [-0.295, -0.235]). After back-transformation, this coefficient
corresponds to an estimated 23.3% reduction in reaction time relative to
the neutral condition (approximately 20.9% to 25.5% lower).

The positive emotional condition was also associated with a significantly
lower log reaction time compared with the neutral condition
(beta = -0.285, SE = 0.015, z = -18.60, p < .001,
95% CI [-0.316, -0.255]). This corresponds to an estimated 24.8%
reduction in reaction time relative to neutral
(approximately 22.5% to 27.1% lower).

The magnitude of the positive and negative coefficients was similar,
suggesting that both emotionally induced conditions were associated with
faster responses compared with the neutral baseline.
""".strip()


INTERPRETATION_TEXT = """
Hypothesis 1 was partially supported.

The emotional manipulation did not produce statistically detectable
differences in response accuracy. Accuracy remained approximately 84-85%
across the three emotional conditions, and the global test of emotional
condition was not significant.

However, emotional condition showed a pronounced association with
reaction speed. Both positive and negative emotional conditions were
associated with substantially faster responses than the neutral condition,
with estimated reductions in reaction time of approximately 25%.

Importantly, the faster responses under emotional induction were not
accompanied by a statistically significant loss of accuracy. Thus, the
results suggest that emotional induction influenced the speed component
of working-memory performance more strongly than the probability of a
correct response.

Because both positive and negative conditions showed effects in the same
direction relative to neutral, these results do not yet demonstrate a
valence-specific advantage of positive emotion or a valence-specific
impairment under negative emotion. Those directional predictions should
be evaluated through the planned contrasts in Hypotheses 2 and 3.
""".strip()


# ======================================================================
# DOCX HELPERS
# ======================================================================

def configure_docx_styles(
    document: Document,
) -> None:

    styles = document.styles

    normal = styles["Normal"]

    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)

    styles["Title"].font.name = "Arial"
    styles["Title"].font.size = Pt(18)
    styles["Title"].font.bold = True

    styles["Heading 1"].font.name = "Arial"
    styles["Heading 1"].font.size = Pt(14)
    styles["Heading 1"].font.bold = True

    styles["Heading 2"].font.name = "Arial"
    styles["Heading 2"].font.size = Pt(12)
    styles["Heading 2"].font.bold = True


def add_dataframe_to_docx(
    document: Document,
    dataframe: pd.DataFrame,
) -> None:

    table = document.add_table(
        rows=1,
        cols=len(dataframe.columns),
    )

    table.alignment = (
        WD_TABLE_ALIGNMENT.CENTER
    )

    table.style = "Table Grid"

    header = table.rows[0].cells

    for i, column in enumerate(
        dataframe.columns
    ):
        header[i].text = str(column)

    for _, row in dataframe.iterrows():

        cells = table.add_row().cells

        for i, value in enumerate(
            row
        ):
            cells[i].text = str(value)


# ======================================================================
# BUILD DOCX
# ======================================================================

def create_docx_report() -> None:

    document = Document()

    configure_docx_styles(
        document
    )

    section = document.sections[0]

    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    # --------------------------------------------------------------
    # Title
    # --------------------------------------------------------------

    title = document.add_heading(
        REPORT_TITLE,
        level=0,
    )

    title.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    subtitle = document.add_paragraph(
        "Confirmatory Hypothesis Report"
    )

    subtitle.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    document.add_paragraph()

    # --------------------------------------------------------------
    # Hypothesis
    # --------------------------------------------------------------

    document.add_heading(
        "Hypothesis",
        level=1,
    )

    document.add_paragraph(
        HYPOTHESIS
    )

    # --------------------------------------------------------------
    # Method
    # --------------------------------------------------------------

    document.add_heading(
        "Method",
        level=1,
    )

    for paragraph_text in (
        METHOD_TEXT.split("\n\n")
    ):

        document.add_paragraph(
            paragraph_text
        )

    # --------------------------------------------------------------
    # Results
    # --------------------------------------------------------------

    document.add_heading(
        "Results",
        level=1,
    )

    document.add_heading(
        "Accuracy",
        level=2,
    )

    for paragraph_text in (
        ACCURACY_RESULTS_TEXT.split(
            "\n\n"
        )
    ):

        document.add_paragraph(
            paragraph_text
        )

    # Accuracy descriptive table

    accuracy_table = (
        accuracy_descriptive.copy()
    )

    accuracy_table["Accuracy"] = (
        accuracy_table["Accuracy"]
        .map(
            lambda x: f"{x * 100:.1f}%"
        )
    )

    accuracy_table["SD"] = (
        accuracy_table["SD"]
        .map(
            lambda x: f"{x:.3f}"
        )
    )

    document.add_paragraph(
        "Table 1. Descriptive accuracy "
        "by emotional condition."
    )

    add_dataframe_to_docx(
        document,
        accuracy_table,
    )

    document.add_paragraph()

    # Accuracy model

    acc_model_table = (
        accuracy_model[
            [
                "Contrast",
                "Odds_Ratio",
                "OR_CI_low",
                "OR_CI_high",
                "p",
            ]
        ]
        .copy()
    )

    acc_model_table.columns = [
        "Contrast",
        "OR",
        "95% CI lower",
        "95% CI upper",
        "p",
    ]

    acc_model_table["OR"] = (
        acc_model_table["OR"]
        .map(
            lambda x: f"{x:.2f}"
        )
    )

    acc_model_table[
        "95% CI lower"
    ] = (
        acc_model_table[
            "95% CI lower"
        ]
        .map(
            lambda x: f"{x:.2f}"
        )
    )

    acc_model_table[
        "95% CI upper"
    ] = (
        acc_model_table[
            "95% CI upper"
        ]
        .map(
            lambda x: f"{x:.2f}"
        )
    )

    acc_model_table["p"] = (
        acc_model_table["p"]
        .map(format_p)
    )

    document.add_paragraph(
        "Table 2. GEE estimates for "
        "response accuracy."
    )

    add_dataframe_to_docx(
        document,
        acc_model_table,
    )

    # --------------------------------------------------------------
    # RT
    # --------------------------------------------------------------

    document.add_heading(
        "Reaction Time",
        level=2,
    )

    for paragraph_text in (
        RT_RESULTS_TEXT.split(
            "\n\n"
        )
    ):

        document.add_paragraph(
            paragraph_text
        )

    rt_table = (
        rt_model[
            [
                "Contrast",
                "Beta_log_RT",
                "Percent_change",
                "Percent_CI_low",
                "Percent_CI_high",
                "p",
            ]
        ]
        .copy()
    )

    rt_table.columns = [
        "Contrast",
        "Beta",
        "RT change",
        "95% CI lower",
        "95% CI upper",
        "p",
    ]

    rt_table["Beta"] = (
        rt_table["Beta"]
        .map(
            lambda x: f"{x:.3f}"
        )
    )

    rt_table["RT change"] = (
        rt_table["RT change"]
        .map(
            lambda x: f"{x:.1f}%"
        )
    )

    rt_table[
        "95% CI lower"
    ] = (
        rt_table[
            "95% CI lower"
        ]
        .map(
            lambda x: f"{x:.1f}%"
        )
    )

    rt_table[
        "95% CI upper"
    ] = (
        rt_table[
            "95% CI upper"
        ]
        .map(
            lambda x: f"{x:.1f}%"
        )
    )

    rt_table["p"] = (
        rt_table["p"]
        .map(format_p)
    )

    document.add_paragraph(
        "Table 3. Mixed-effects model "
        "estimates for log reaction time."
    )

    add_dataframe_to_docx(
        document,
        rt_table,
    )

    # --------------------------------------------------------------
    # Interpretation
    # --------------------------------------------------------------

    document.add_heading(
        "Interpretation",
        level=1,
    )

    for paragraph_text in (
        INTERPRETATION_TEXT.split(
            "\n\n"
        )
    ):

        document.add_paragraph(
            paragraph_text
        )

    document.add_heading(
        "Hypothesis Assessment",
        level=1,
    )

    paragraph = document.add_paragraph()

    run = paragraph.add_run(
        "Conclusion: "
    )

    run.bold = True

    paragraph.add_run(
        "Hypothesis 1 was partially supported. "
        "Emotional condition did not significantly "
        "affect accuracy, but both positive and "
        "negative emotional conditions were associated "
        "with substantially faster reaction times "
        "relative to the neutral condition."
    )

    # --------------------------------------------------------------
    # Save
    # --------------------------------------------------------------

    document.save(
        DOCX_PATH
    )

    print(
        f"DOCX created: {DOCX_PATH}"
    )


# ======================================================================
# PDF HELPERS
# ======================================================================

def dataframe_to_pdf_table(
    dataframe: pd.DataFrame,
    col_widths=None,
):

    data = [
        list(dataframe.columns)
    ]

    for _, row in dataframe.iterrows():

        data.append(
            [
                str(value)
                for value in row
            ]
        )

    table = Table(
        data,
        colWidths=col_widths,
        repeatRows=1,
    )

    table.setStyle(
        TableStyle(
            [
                (
                    "BACKGROUND",
                    (0, 0),
                    (-1, 0),
                    colors.lightgrey,
                ),
                (
                    "TEXTCOLOR",
                    (0, 0),
                    (-1, 0),
                    colors.black,
                ),
                (
                    "FONTNAME",
                    (0, 0),
                    (-1, 0),
                    "Helvetica-Bold",
                ),
                (
                    "FONTNAME",
                    (0, 1),
                    (-1, -1),
                    "Helvetica",
                ),
                (
                    "FONTSIZE",
                    (0, 0),
                    (-1, -1),
                    8,
                ),
                (
                    "GRID",
                    (0, 0),
                    (-1, -1),
                    0.5,
                    colors.grey,
                ),
                (
                    "VALIGN",
                    (0, 0),
                    (-1, -1),
                    "MIDDLE",
                ),
                (
                    "LEFTPADDING",
                    (0, 0),
                    (-1, -1),
                    5,
                ),
                (
                    "RIGHTPADDING",
                    (0, 0),
                    (-1, -1),
                    5,
                ),
                (
                    "TOPPADDING",
                    (0, 0),
                    (-1, -1),
                    5,
                ),
                (
                    "BOTTOMPADDING",
                    (0, 0),
                    (-1, -1),
                    5,
                ),
            ]
        )
    )

    return table


# ======================================================================
# BUILD PDF
# ======================================================================

def create_pdf_report() -> None:

    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        "ReportTitle",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=17,
        leading=21,
        alignment=TA_CENTER,
        spaceAfter=18,
    )

    heading1 = ParagraphStyle(
        "Heading1Custom",
        parent=styles["Heading1"],
        fontName="Helvetica-Bold",
        fontSize=13,
        leading=16,
        spaceBefore=10,
        spaceAfter=7,
    )

    heading2 = ParagraphStyle(
        "Heading2Custom",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=11,
        leading=14,
        spaceBefore=8,
        spaceAfter=6,
    )

    body_style = ParagraphStyle(
        "BodyCustom",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=9.5,
        leading=14,
        spaceAfter=8,
    )

    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=A4,
        rightMargin=1.8 * cm,
        leftMargin=1.8 * cm,
        topMargin=1.8 * cm,
        bottomMargin=1.8 * cm,
        title=REPORT_TITLE,
    )

    story = []

    # --------------------------------------------------------------
    # Title
    # --------------------------------------------------------------

    story.append(
        Paragraph(
            REPORT_TITLE,
            title_style,
        )
    )

    story.append(
        Paragraph(
            "Confirmatory Hypothesis Report",
            body_style,
        )
    )

    story.append(
        Spacer(
            1,
            10,
        )
    )

    # --------------------------------------------------------------
    # Hypothesis
    # --------------------------------------------------------------

    story.append(
        Paragraph(
            "Hypothesis",
            heading1,
        )
    )

    story.append(
        Paragraph(
            HYPOTHESIS,
            body_style,
        )
    )

    # --------------------------------------------------------------
    # Method
    # --------------------------------------------------------------

    story.append(
        Paragraph(
            "Method",
            heading1,
        )
    )

    for text in METHOD_TEXT.split(
        "\n\n"
    ):

        story.append(
            Paragraph(
                text,
                body_style,
            )
        )

    # --------------------------------------------------------------
    # Results - Accuracy
    # --------------------------------------------------------------

    story.append(
        Paragraph(
            "Results",
            heading1,
        )
    )

    story.append(
        Paragraph(
            "Accuracy",
            heading2,
        )
    )

    for text in (
        ACCURACY_RESULTS_TEXT
        .split("\n\n")
    ):

        story.append(
            Paragraph(
                text,
                body_style,
            )
        )

    accuracy_table = (
        accuracy_descriptive.copy()
    )

    accuracy_table[
        "Accuracy"
    ] = (
        accuracy_table[
            "Accuracy"
        ]
        .map(
            lambda x:
            f"{x * 100:.1f}%"
        )
    )

    accuracy_table[
        "SD"
    ] = (
        accuracy_table[
            "SD"
        ]
        .map(
            lambda x:
            f"{x:.3f}"
        )
    )

    story.append(
        Paragraph(
            "Table 1. Descriptive accuracy by emotional condition.",
            body_style,
        )
    )

    story.append(
        dataframe_to_pdf_table(
            accuracy_table,
            col_widths=[
                3.5 * cm,
                2.5 * cm,
                2.5 * cm,
                2.5 * cm,
            ],
        )
    )

    story.append(
        Spacer(
            1,
            12,
        )
    )

    acc_model_table = (
        accuracy_model[
            [
                "Contrast",
                "Odds_Ratio",
                "OR_CI_low",
                "OR_CI_high",
                "p",
            ]
        ]
        .copy()
    )

    acc_model_table.columns = [
        "Contrast",
        "OR",
        "CI low",
        "CI high",
        "p",
    ]

    acc_model_table[
        "OR"
    ] = (
        acc_model_table[
            "OR"
        ]
        .map(
            lambda x:
            f"{x:.2f}"
        )
    )

    acc_model_table[
        "CI low"
    ] = (
        acc_model_table[
            "CI low"
        ]
        .map(
            lambda x:
            f"{x:.2f}"
        )
    )

    acc_model_table[
        "CI high"
    ] = (
        acc_model_table[
            "CI high"
        ]
        .map(
            lambda x:
            f"{x:.2f}"
        )
    )

    acc_model_table[
        "p"
    ] = (
        acc_model_table[
            "p"
        ]
        .map(format_p)
    )

    story.append(
        Paragraph(
            "Table 2. GEE estimates for response accuracy.",
            body_style,
        )
    )

    story.append(
        dataframe_to_pdf_table(
            acc_model_table,
            col_widths=[
                5.0 * cm,
                2.0 * cm,
                2.2 * cm,
                2.2 * cm,
                2.0 * cm,
            ],
        )
    )

    story.append(
        Spacer(
            1,
            15,
        )
    )

    # --------------------------------------------------------------
    # RT
    # --------------------------------------------------------------

    story.append(
        Paragraph(
            "Reaction Time",
            heading2,
        )
    )

    for text in (
        RT_RESULTS_TEXT
        .split("\n\n")
    ):

        story.append(
            Paragraph(
                text,
                body_style,
            )
        )

    rt_table = (
        rt_model[
            [
                "Contrast",
                "Beta_log_RT",
                "Percent_change",
                "Percent_CI_low",
                "Percent_CI_high",
                "p",
            ]
        ]
        .copy()
    )

    rt_table.columns = [
        "Contrast",
        "Beta",
        "RT change",
        "CI low",
        "CI high",
        "p",
    ]

    rt_table[
        "Beta"
    ] = (
        rt_table[
            "Beta"
        ]
        .map(
            lambda x:
            f"{x:.3f}"
        )
    )

    for column in [
        "RT change",
        "CI low",
        "CI high",
    ]:

        rt_table[
            column
        ] = (
            rt_table[
                column
            ]
            .map(
                lambda x:
                f"{x:.1f}%"
            )
        )

    rt_table[
        "p"
    ] = (
        rt_table[
            "p"
        ]
        .map(format_p)
    )

    story.append(
        Paragraph(
            "Table 3. Mixed-effects model estimates for log reaction time.",
            body_style,
        )
    )

    story.append(
        dataframe_to_pdf_table(
            rt_table,
            col_widths=[
                4.0 * cm,
                1.7 * cm,
                2.2 * cm,
                2.0 * cm,
                2.0 * cm,
                1.6 * cm,
            ],
        )
    )

    # --------------------------------------------------------------
    # Interpretation
    # --------------------------------------------------------------

    story.append(
        Spacer(
            1,
            15,
        )
    )

    story.append(
        Paragraph(
            "Interpretation",
            heading1,
        )
    )

    for text in (
        INTERPRETATION_TEXT
        .split("\n\n")
    ):

        story.append(
            Paragraph(
                text,
                body_style,
            )
        )

    story.append(
        Paragraph(
            "Hypothesis Assessment",
            heading1,
        )
    )

    conclusion = (
        "<b>Conclusion:</b> "
        "Hypothesis 1 was partially supported. "
        "Emotional condition did not significantly "
        "affect accuracy, but both positive and negative "
        "emotional conditions were associated with "
        "substantially faster reaction times relative "
        "to the neutral condition."
    )

    story.append(
        Paragraph(
            conclusion,
            body_style,
        )
    )

    doc.build(
        story
    )

    print(
        f"PDF created: {PDF_PATH}"
    )


# ======================================================================
# MAIN
# ======================================================================

def main():

    print(
        "=" * 70
    )

    print(
        "GENERATING HYPOTHESIS 1 REPORT"
    )

    print(
        "=" * 70
    )

    create_docx_report()

    create_pdf_report()

    print(
        "\nReport generation completed."
    )

    print(
        f"\nDOCX:\n{DOCX_PATH}"
    )

    print(
        f"\nPDF:\n{PDF_PATH}"
    )


if __name__ == "__main__":
    main()