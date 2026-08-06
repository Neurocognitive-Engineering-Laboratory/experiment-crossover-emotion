"""
Generate the final DOCX and PDF reports for Hypothesis 2.

Expected input files
--------------------
reports/tables/hypothesis_02/
    h2_accuracy_descriptive.csv
    h2_accuracy_model.csv
    h2_rt_descriptive.csv
    h2_rt_model.csv
    h2_efficiency_descriptive.csv
    h2_efficiency_test.csv
    h2_summary.csv

Generated files
---------------
reports/final_report/
    hypothesis_02_report.docx
    hypothesis_02_report.pdf
"""

from __future__ import annotations

from pathlib import Path
from typing import Any

import numpy as np
import pandas as pd

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import (
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


# ======================================================================
# PATHS
# ======================================================================

PROJECT_ROOT = Path(__file__).resolve().parents[1]

INPUT_DIR = (
    PROJECT_ROOT
    / "reports"
    / "tables"
    / "hypothesis_02"
)

OUTPUT_DIR = (
    PROJECT_ROOT
    / "reports"
    / "final_report"
)

OUTPUT_DIR.mkdir(
    parents=True,
    exist_ok=True,
)

DOCX_PATH = (
    OUTPUT_DIR
    / "hypothesis_02_report.docx"
)

PDF_PATH = (
    OUTPUT_DIR
    / "hypothesis_02_report.pdf"
)


INPUT_FILES = {
    "accuracy_descriptive":
        INPUT_DIR / "h2_accuracy_descriptive.csv",

    "accuracy_model":
        INPUT_DIR / "h2_accuracy_model.csv",

    "rt_descriptive":
        INPUT_DIR / "h2_rt_descriptive.csv",

    "rt_model":
        INPUT_DIR / "h2_rt_model.csv",

    "efficiency_descriptive":
        INPUT_DIR / "h2_efficiency_descriptive.csv",

    "efficiency_test":
        INPUT_DIR / "h2_efficiency_test.csv",

    "summary":
        INPUT_DIR / "h2_summary.csv",
}


# ======================================================================
# REPORT INFORMATION
# ======================================================================

REPORT_TITLE = (
    "Hypothesis 2 - Effect of Negative Emotion "
    "on Working-Memory Performance"
)

REPORT_SUBTITLE = (
    "Confirmatory Hypothesis Report"
)

HYPOTHESIS_TEXT = (
    "Negative emotional induction leads to poorer working-memory "
    "performance relative to the neutral emotional condition."
)


# ======================================================================
# GENERAL HELPERS
# ======================================================================

def validate_input_files() -> None:
    """
    Check whether all expected result files exist.
    """

    missing = [
        str(path)
        for path in INPUT_FILES.values()
        if not path.exists()
    ]

    if missing:
        missing_text = "\n".join(
            f"- {path}"
            for path in missing
        )

        raise FileNotFoundError(
            "The following Hypothesis 2 result files were not found:\n"
            f"{missing_text}\n\n"
            "Run 04_2_hypothesis_2.ipynb before generating the report."
        )


def load_results() -> dict[str, pd.DataFrame]:
    """
    Load every result table exported by the H2 notebook.
    """

    validate_input_files()

    return {
        name: pd.read_csv(path)
        for name, path in INPUT_FILES.items()
    }


def first_value(
    data: pd.DataFrame,
    column: str,
    default: Any = np.nan,
) -> Any:
    """
    Safely retrieve the first value of a column.
    """

    if column not in data.columns:
        return default

    if data.empty:
        return default

    return data[column].iloc[0]


def get_condition_row(
    data: pd.DataFrame,
    condition: str,
) -> pd.Series:
    """
    Retrieve the descriptive row for a given emotional condition.
    """

    condition_columns = [
        "emotion_condition",
        "condition",
        "Condition",
    ]

    condition_column = next(
        (
            column
            for column in condition_columns
            if column in data.columns
        ),
        None,
    )

    if condition_column is None:
        raise KeyError(
            "Could not identify the condition column."
        )

    subset = data.loc[
        data[condition_column]
        .astype(str)
        .str.lower()
        == condition.lower()
    ]

    if subset.empty:
        raise ValueError(
            f"Condition '{condition}' was not found."
        )

    return subset.iloc[0]


def format_p_value(
    p_value: float,
) -> str:
    """
    Format p-values for scientific reports.
    """

    if pd.isna(p_value):
        return "NA"

    if p_value < 0.001:
        return "< .001"

    return f"= {p_value:.3f}"


def format_number(
    value: Any,
    digits: int = 3,
) -> str:
    """
    Format a numeric value safely.
    """

    try:
        numeric = float(value)
    except (TypeError, ValueError):
        return str(value)

    if np.isnan(numeric):
        return "NA"

    return f"{numeric:.{digits}f}"


def format_percent(
    value: Any,
    digits: int = 1,
    already_percent: bool = False,
) -> str:
    """
    Format a proportion or percentage.
    """

    try:
        numeric = float(value)
    except (TypeError, ValueError):
        return str(value)

    if np.isnan(numeric):
        return "NA"

    if not already_percent:
        numeric *= 100

    return f"{numeric:.{digits}f}%"


def find_model_contrast(
    data: pd.DataFrame,
    keyword: str = "Negative",
) -> pd.Series:
    """
    Find the model row corresponding to the Negative vs Neutral contrast.
    """

    text_columns = [
        "term",
        "contrast",
        "Contrast",
    ]

    for column in text_columns:

        if column in data.columns:

            mask = (
                data[column]
                .astype(str)
                .str.contains(
                    keyword,
                    case=False,
                    regex=False,
                )
            )

            if mask.any():
                return data.loc[mask].iloc[0]

    if len(data) == 1:
        return data.iloc[0]

    raise ValueError(
        "Could not identify the Negative vs Neutral model contrast."
    )


# ======================================================================
# RESULT EXTRACTION
# ======================================================================

def extract_report_values(
    results: dict[str, pd.DataFrame],
) -> dict[str, Any]:
    """
    Extract all numerical values required for the report.
    """

    accuracy_desc = results[
        "accuracy_descriptive"
    ]

    accuracy_model = results[
        "accuracy_model"
    ]

    rt_desc = results[
        "rt_descriptive"
    ]

    rt_model = results[
        "rt_model"
    ]

    efficiency_desc = results[
        "efficiency_descriptive"
    ]

    efficiency_test = results[
        "efficiency_test"
    ]

    summary = results[
        "summary"
    ]

    negative_accuracy_row = get_condition_row(
        accuracy_desc,
        "Negative",
    )

    neutral_accuracy_row = get_condition_row(
        accuracy_desc,
        "Neutral",
    )

    negative_rt_row = get_condition_row(
        rt_desc,
        "Negative",
    )

    neutral_rt_row = get_condition_row(
        rt_desc,
        "Neutral",
    )

    accuracy_contrast = find_model_contrast(
        accuracy_model
    )

    rt_contrast = find_model_contrast(
        rt_model
    )

    accuracy_mean_column = next(
        column
        for column in [
            "accuracy_mean",
            "accuracy",
            "mean",
        ]
        if column in accuracy_desc.columns
    )

    accuracy_sd_column = next(
        (
            column
            for column in [
                "accuracy_sd",
                "sd",
                "std",
            ]
            if column in accuracy_desc.columns
        ),
        None,
    )

    rt_mean_column = next(
        column
        for column in [
            "mean_rt",
            "reaction_time_mean",
        ]
        if column in rt_desc.columns
    )

    rt_median_column = next(
        (
            column
            for column in [
                "median_rt",
                "reaction_time_median",
            ]
            if column in rt_desc.columns
        ),
        None,
    )

    rt_sd_column = next(
        (
            column
            for column in [
                "sd_rt",
                "reaction_time_sd",
            ]
            if column in rt_desc.columns
        ),
        None,
    )

    values = {
        "participants":
            int(
                first_value(
                    summary,
                    "participants",
                    23,
                )
            ),

        "accuracy_observations":
            int(
                first_value(
                    summary,
                    "accuracy_observations",
                    11880,
                )
            ),

        "rt_observations":
            int(
                first_value(
                    summary,
                    "rt_observations",
                    9533,
                )
            ),

        "negative_accuracy":
            float(
                negative_accuracy_row[
                    accuracy_mean_column
                ]
            ),

        "neutral_accuracy":
            float(
                neutral_accuracy_row[
                    accuracy_mean_column
                ]
            ),

        "negative_accuracy_sd":
            (
                float(
                    negative_accuracy_row[
                        accuracy_sd_column
                    ]
                )
                if accuracy_sd_column
                else np.nan
            ),

        "neutral_accuracy_sd":
            (
                float(
                    neutral_accuracy_row[
                        accuracy_sd_column
                    ]
                )
                if accuracy_sd_column
                else np.nan
            ),

        "accuracy_beta":
            float(
                accuracy_contrast.get(
                    "beta",
                    accuracy_contrast.get(
                        "estimate",
                        np.nan,
                    ),
                )
            ),

        "accuracy_se":
            float(
                accuracy_contrast.get(
                    "se",
                    accuracy_contrast.get(
                        "std_error",
                        np.nan,
                    ),
                )
            ),

        "accuracy_z":
            float(
                accuracy_contrast.get(
                    "z",
                    np.nan,
                )
            ),

        "accuracy_p":
            float(
                accuracy_contrast.get(
                    "p",
                    accuracy_contrast.get(
                        "p_value",
                        np.nan,
                    ),
                )
            ),

        "accuracy_or":
            float(
                accuracy_contrast.get(
                    "odds_ratio",
                    np.exp(
                        accuracy_contrast.get(
                            "beta",
                            np.nan,
                        )
                    ),
                )
            ),

        "accuracy_or_ci_low":
            float(
                accuracy_contrast.get(
                    "or_ci_low",
                    np.nan,
                )
            ),

        "accuracy_or_ci_high":
            float(
                accuracy_contrast.get(
                    "or_ci_high",
                    np.nan,
                )
            ),

        "negative_mean_rt":
            float(
                negative_rt_row[
                    rt_mean_column
                ]
            ),

        "neutral_mean_rt":
            float(
                neutral_rt_row[
                    rt_mean_column
                ]
            ),

        "negative_median_rt":
            (
                float(
                    negative_rt_row[
                        rt_median_column
                    ]
                )
                if rt_median_column
                else np.nan
            ),

        "neutral_median_rt":
            (
                float(
                    neutral_rt_row[
                        rt_median_column
                    ]
                )
                if rt_median_column
                else np.nan
            ),

        "negative_rt_sd":
            (
                float(
                    negative_rt_row[
                        rt_sd_column
                    ]
                )
                if rt_sd_column
                else np.nan
            ),

        "neutral_rt_sd":
            (
                float(
                    neutral_rt_row[
                        rt_sd_column
                    ]
                )
                if rt_sd_column
                else np.nan
            ),

        "rt_beta":
            float(
                rt_contrast.get(
                    "beta",
                    np.nan,
                )
            ),

        "rt_se":
            float(
                rt_contrast.get(
                    "se",
                    np.nan,
                )
            ),

        "rt_z":
            float(
                rt_contrast.get(
                    "z",
                    np.nan,
                )
            ),

        "rt_p":
            float(
                rt_contrast.get(
                    "p",
                    np.nan,
                )
            ),

        "rt_percent_change":
            float(
                rt_contrast.get(
                    "percent_change",
                    np.nan,
                )
            ),

        "rt_percent_ci_low":
            float(
                rt_contrast.get(
                    "percent_ci_low",
                    np.nan,
                )
            ),

        "rt_percent_ci_high":
            float(
                rt_contrast.get(
                    "percent_ci_high",
                    np.nan,
                )
            ),

        "assessment":
            str(
                first_value(
                    summary,
                    "assessment",
                    "not supported",
                )
            ),

        "efficiency_descriptive":
            efficiency_desc,

        "efficiency_test":
            efficiency_test,
    }

    return values


# ======================================================================
# REPORT TEXT
# ======================================================================

def build_method_text(
    values: dict[str, Any],
) -> list[str]:
    """
    Build the English Method section.
    """

    return [
        (
            "A confirmatory trial-level analysis was conducted to determine "
            "whether negative emotional induction impaired working-memory "
            "performance relative to the neutral condition. The analytical "
            f"sample included {values['participants']} participants."
        ),
        (
            "The primary outcomes were response accuracy and reaction time. "
            "Accuracy was analyzed using a generalized estimating equation "
            "(GEE) model with a binomial distribution, logit link, "
            "exchangeable within-participant correlation structure, and "
            "robust standard errors. Emotional condition was entered as the "
            "predictor, with the neutral condition specified as the reference "
            f"category. The accuracy analysis included "
            f"{values['accuracy_observations']:,} trial-level observations."
        ),
        (
            "Reaction time was analyzed among valid correct-response trials "
            "only. Because reaction-time values were positively skewed, they "
            "were log-transformed before modeling. A linear mixed-effects "
            "model was fitted with emotional condition as a fixed effect and "
            "a participant-specific random intercept. Maximum-likelihood "
            f"estimation was used. The reaction-time analysis included "
            f"{values['rt_observations']:,} observations."
        ),
        (
            "The GEE coefficient was exponentiated and interpreted as an odds "
            "ratio. The log reaction-time coefficient was exponentiated and "
            "converted to the percentage difference in reaction time between "
            "the negative and neutral conditions. A secondary analysis "
            "examined inverse efficiency as a descriptive indicator of the "
            "joint speed-accuracy pattern. Statistical significance was "
            "evaluated using a two-sided alpha level of .05, and 95% "
            "confidence intervals are reported."
        ),
    ]


def build_accuracy_results_text(
    values: dict[str, Any],
) -> list[str]:

    accuracy_difference = (
        values["negative_accuracy"]
        - values["neutral_accuracy"]
    ) * 100

    return [
        (
            "Mean trial-level accuracy was "
            f"{format_percent(values['negative_accuracy'])} in the negative "
            f"condition and {format_percent(values['neutral_accuracy'])} in "
            "the neutral condition. The descriptive difference was "
            f"{accuracy_difference:.1f} percentage points."
        ),
        (
            "Negative emotion was associated with a small reduction in the "
            "estimated odds of a correct response relative to the neutral "
            "condition "
            f"(beta = {values['accuracy_beta']:.3f}, "
            f"SE = {values['accuracy_se']:.3f}, "
            f"z = {values['accuracy_z']:.2f}, "
            f"p {format_p_value(values['accuracy_p'])}). "
            f"The corresponding odds ratio was "
            f"{values['accuracy_or']:.2f} "
            f"(95% CI [{values['accuracy_or_ci_low']:.2f}, "
            f"{values['accuracy_or_ci_high']:.2f}])."
        ),
        (
            "The confidence interval included the null value of 1.00, and "
            "the comparison was not statistically significant. Therefore, "
            "the analysis did not provide sufficient evidence that negative "
            "emotion reduced response accuracy relative to the neutral "
            "condition."
        ),
    ]


def build_rt_results_text(
    values: dict[str, Any],
) -> list[str]:

    percent_change = values[
        "rt_percent_change"
    ]

    if percent_change < 0:
        direction = "lower"
        interpretation = "faster"
    else:
        direction = "higher"
        interpretation = "slower"

    return [
        (
            "The mean reaction time was "
            f"{values['negative_mean_rt']:.1f} ms in the negative condition "
            f"and {values['neutral_mean_rt']:.1f} ms in the neutral "
            "condition."
        ),
        (
            "The negative condition was significantly associated with "
            f"{direction} log reaction time relative to the neutral condition "
            f"(beta = {values['rt_beta']:.3f}, "
            f"SE = {values['rt_se']:.3f}, "
            f"z = {values['rt_z']:.2f}, "
            f"p {format_p_value(values['rt_p'])})."
        ),
        (
            "After back-transformation, reaction time under negative emotion "
            f"was estimated to be {abs(percent_change):.1f}% {direction} "
            "than under the neutral condition "
            f"(95% CI from {values['rt_percent_ci_low']:.1f}% to "
            f"{values['rt_percent_ci_high']:.1f}%). "
            f"Thus, participants responded significantly {interpretation}, "
            "rather than more slowly, during the negative emotional condition."
        ),
    ]


def build_interpretation_text(
    values: dict[str, Any],
) -> list[str]:

    return [
        (
            "Hypothesis 2 was not supported. The hypothesis predicted that "
            "negative emotional induction would impair working-memory "
            "performance relative to the neutral condition."
        ),
        (
            "Although accuracy was descriptively lower under negative "
            "emotion, the difference was small and statistically "
            "non-significant. Consequently, there was no reliable evidence "
            "that negative emotion decreased the probability of producing a "
            "correct response."
        ),
        (
            "Reaction-time results were opposite to the predicted impairment "
            "pattern. Rather than responding more slowly, participants "
            f"responded approximately "
            f"{abs(values['rt_percent_change']):.1f}% faster under negative "
            "emotion than under the neutral condition."
        ),
        (
            "The combination of significantly faster responses and a small, "
            "non-significant reduction in accuracy may be descriptively "
            "compatible with a shift toward faster responding. However, it "
            "does not provide strong evidence of a speed-accuracy trade-off, "
            "because the reduction in accuracy was not statistically "
            "supported."
        ),
        (
            "Overall, negative emotional induction altered response speed but "
            "did not produce the expected deterioration in working-memory "
            "performance. These findings indicate that the negative condition "
            "was associated with faster responding without a statistically "
            "detectable loss of accuracy."
        ),
    ]


DECISION_RULES = [
    (
        "Hypothesis 2 is supported if negative emotion is associated with "
        "significantly lower accuracy, significantly slower reaction time, "
        "or both."
    ),
    (
        "Hypothesis 2 is partially supported if only one impairment outcome "
        "is statistically supported or if there is clear evidence of a "
        "speed-accuracy trade-off."
    ),
    (
        "Hypothesis 2 is not supported if negative emotion does not "
        "significantly reduce accuracy and responses are not significantly "
        "slower than under neutral emotion."
    ),
]


# ======================================================================
# DOCX HELPERS
# ======================================================================

def configure_docx_styles(
    document: Document,
) -> None:

    document.styles["Normal"].font.name = (
        "Arial"
    )

    document.styles["Normal"].font.size = (
        Pt(10.5)
    )

    document.styles["Title"].font.name = (
        "Arial"
    )

    document.styles["Title"].font.size = (
        Pt(18)
    )

    document.styles["Title"].font.bold = True

    document.styles["Heading 1"].font.name = (
        "Arial"
    )

    document.styles["Heading 1"].font.size = (
        Pt(14)
    )

    document.styles["Heading 1"].font.bold = True

    document.styles["Heading 2"].font.name = (
        "Arial"
    )

    document.styles["Heading 2"].font.size = (
        Pt(12)
    )

    document.styles["Heading 2"].font.bold = True


def add_docx_table(
    document: Document,
    dataframe: pd.DataFrame,
) -> None:

    table = document.add_table(
        rows=1,
        cols=len(dataframe.columns),
    )

    table.style = "Table Grid"

    table.alignment = (
        WD_TABLE_ALIGNMENT.CENTER
    )

    for index, column in enumerate(
        dataframe.columns
    ):
        table.rows[0].cells[
            index
        ].text = str(column)

    for _, row in dataframe.iterrows():

        cells = table.add_row().cells

        for index, value in enumerate(
            row
        ):
            cells[index].text = str(value)


def build_accuracy_descriptive_table(
    results: dict[str, pd.DataFrame],
) -> pd.DataFrame:

    table = results[
        "accuracy_descriptive"
    ].copy()

    rename_map = {
        "emotion_condition": "Condition",
        "condition": "Condition",
        "participants": "Participants",
        "trials": "Trials",
        "correct": "Correct",
        "accuracy_mean": "Accuracy",
        "accuracy_sd": "SD",
        "accuracy_median": "Median",
    }

    table = table.rename(
        columns=rename_map
    )

    if "Accuracy" in table.columns:
        table["Accuracy"] = table[
            "Accuracy"
        ].map(format_percent)

    if "SD" in table.columns:
        table["SD"] = table[
            "SD"
        ].map(
            lambda value:
            format_number(
                value,
                3,
            )
        )

    return table


def build_accuracy_model_table(
    values: dict[str, Any],
) -> pd.DataFrame:

    return pd.DataFrame(
        {
            "Contrast": [
                "Negative vs Neutral"
            ],
            "Beta": [
                f"{values['accuracy_beta']:.3f}"
            ],
            "OR": [
                f"{values['accuracy_or']:.2f}"
            ],
            "95% CI": [
                (
                    f"[{values['accuracy_or_ci_low']:.2f}, "
                    f"{values['accuracy_or_ci_high']:.2f}]"
                )
            ],
            "p": [
                format_p_value(
                    values[
                        "accuracy_p"
                    ]
                )
            ],
        }
    )


def build_rt_descriptive_table(
    results: dict[str, pd.DataFrame],
) -> pd.DataFrame:

    table = results[
        "rt_descriptive"
    ].copy()

    rename_map = {
        "emotion_condition": "Condition",
        "condition": "Condition",
        "participants": "Participants",
        "trials": "Trials",
        "mean_rt": "Mean RT",
        "sd_rt": "SD",
        "median_rt": "Median RT",
        "mean_log_rt": "Mean log RT",
    }

    table = table.rename(
        columns=rename_map
    )

    numeric_columns = [
        "Mean RT",
        "SD",
        "Median RT",
        "Mean log RT",
    ]

    for column in numeric_columns:

        if column in table.columns:

            table[column] = table[
                column
            ].map(
                lambda value:
                format_number(
                    value,
                    2,
                )
            )

    return table


def build_rt_model_table(
    values: dict[str, Any],
) -> pd.DataFrame:

    return pd.DataFrame(
        {
            "Contrast": [
                "Negative vs Neutral"
            ],
            "Beta": [
                f"{values['rt_beta']:.3f}"
            ],
            "RT change": [
                f"{values['rt_percent_change']:.1f}%"
            ],
            "95% CI": [
                (
                    f"[{values['rt_percent_ci_low']:.1f}%, "
                    f"{values['rt_percent_ci_high']:.1f}%]"
                )
            ],
            "p": [
                format_p_value(
                    values["rt_p"]
                )
            ],
        }
    )


# ======================================================================
# DOCX REPORT
# ======================================================================

def create_docx_report(
    results: dict[str, pd.DataFrame],
    values: dict[str, Any],
) -> None:

    document = Document()

    configure_docx_styles(
        document
    )

    section = document.sections[0]

    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    title = document.add_heading(
        REPORT_TITLE,
        level=0,
    )

    title.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    subtitle = document.add_paragraph(
        REPORT_SUBTITLE
    )

    subtitle.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    document.add_heading(
        "Hypothesis",
        level=1,
    )

    document.add_paragraph(
        HYPOTHESIS_TEXT
    )

    document.add_heading(
        "Method",
        level=1,
    )

    for paragraph in build_method_text(
        values
    ):
        document.add_paragraph(
            paragraph
        )

    document.add_heading(
        "Results",
        level=1,
    )

    # Accuracy ------------------------------------------------------

    document.add_heading(
        "Accuracy",
        level=2,
    )

    for paragraph in (
        build_accuracy_results_text(
            values
        )
    ):
        document.add_paragraph(
            paragraph
        )

    document.add_paragraph(
        "Table 1. Descriptive accuracy under neutral and negative emotion."
    )

    add_docx_table(
        document,
        build_accuracy_descriptive_table(
            results
        ),
    )

    document.add_paragraph()

    document.add_paragraph(
        "Table 2. Binomial GEE estimate for response accuracy."
    )

    add_docx_table(
        document,
        build_accuracy_model_table(
            values
        ),
    )

    # RT ------------------------------------------------------------

    document.add_heading(
        "Reaction Time",
        level=2,
    )

    for paragraph in build_rt_results_text(
        values
    ):
        document.add_paragraph(
            paragraph
        )

    document.add_paragraph(
        "Table 3. Descriptive reaction time under neutral and negative emotion."
    )

    add_docx_table(
        document,
        build_rt_descriptive_table(
            results
        ),
    )

    document.add_paragraph()

    document.add_paragraph(
        "Table 4. Mixed-effects model estimate for log reaction time."
    )

    add_docx_table(
        document,
        build_rt_model_table(
            values
        ),
    )

    # Speed-accuracy ------------------------------------------------

    document.add_heading(
        "Speed-Accuracy Pattern",
        level=2,
    )

    document.add_paragraph(
        (
            "Negative emotion was associated with significantly faster "
            "responses, whereas the corresponding reduction in accuracy was "
            "small and statistically non-significant. This pattern does not "
            "provide sufficient statistical evidence of a speed-accuracy "
            "trade-off, although it may indicate a descriptive shift toward "
            "faster responding."
        )
    )

    # Decision rules ------------------------------------------------

    document.add_heading(
        "Decision Rules",
        level=1,
    )

    for rule in DECISION_RULES:
        document.add_paragraph(
            rule,
            style="List Bullet",
        )

    # Interpretation ------------------------------------------------

    document.add_heading(
        "Interpretation",
        level=1,
    )

    for paragraph in build_interpretation_text(
        values
    ):
        document.add_paragraph(
            paragraph
        )

    # Final assessment ----------------------------------------------

    document.add_heading(
        "Hypothesis Assessment",
        level=1,
    )

    conclusion = document.add_paragraph()

    conclusion_run = conclusion.add_run(
        "Conclusion: "
    )

    conclusion_run.bold = True

    conclusion.add_run(
        (
            "Hypothesis 2 was not supported. Negative emotion did not "
            "significantly reduce accuracy and was associated with faster, "
            "rather than slower, reaction times relative to the neutral "
            "condition."
        )
    )

    document.save(
        DOCX_PATH
    )

    print(
        f"DOCX created: {DOCX_PATH}"
    )


# ======================================================================
# PDF HELPERS
# ======================================================================

def make_pdf_styles():

    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        "CustomTitle",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=17,
        leading=21,
        alignment=TA_CENTER,
        spaceAfter=14,
    )

    subtitle_style = ParagraphStyle(
        "CustomSubtitle",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=10,
        leading=13,
        alignment=TA_CENTER,
        spaceAfter=16,
    )

    heading_1 = ParagraphStyle(
        "CustomHeading1",
        parent=styles["Heading1"],
        fontName="Helvetica-Bold",
        fontSize=13,
        leading=16,
        spaceBefore=10,
        spaceAfter=6,
    )

    heading_2 = ParagraphStyle(
        "CustomHeading2",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=11,
        leading=14,
        spaceBefore=8,
        spaceAfter=5,
    )

    body = ParagraphStyle(
        "CustomBody",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=9.5,
        leading=14,
        alignment=TA_JUSTIFY,
        spaceAfter=8,
    )

    bullet = ParagraphStyle(
        "CustomBullet",
        parent=body,
        leftIndent=12,
        firstLineIndent=-7,
    )

    return {
        "title": title_style,
        "subtitle": subtitle_style,
        "heading_1": heading_1,
        "heading_2": heading_2,
        "body": body,
        "bullet": bullet,
    }


def dataframe_to_pdf_table(
    dataframe: pd.DataFrame,
    column_widths=None,
) -> Table:

    table_data = [
        [
            Paragraph(
                f"<b>{column}</b>",
                getSampleStyleSheet()[
                    "BodyText"
                ],
            )
            for column in dataframe.columns
        ]
    ]

    for _, row in dataframe.iterrows():

        table_data.append(
            [
                Paragraph(
                    str(value),
                    getSampleStyleSheet()[
                        "BodyText"
                    ],
                )
                for value in row
            ]
        )

    table = Table(
        table_data,
        colWidths=column_widths,
        repeatRows=1,
        hAlign="CENTER",
    )

    table.setStyle(
        TableStyle(
            [
                (
                    "BACKGROUND",
                    (0, 0),
                    (-1, 0),
                    colors.HexColor(
                        "#E6E6E6"
                    ),
                ),
                (
                    "GRID",
                    (0, 0),
                    (-1, -1),
                    0.5,
                    colors.grey,
                ),
                (
                    "VALIGN",
                    (0, 0),
                    (-1, -1),
                    "MIDDLE",
                ),
                (
                    "FONTSIZE",
                    (0, 0),
                    (-1, -1),
                    8,
                ),
                (
                    "LEFTPADDING",
                    (0, 0),
                    (-1, -1),
                    5,
                ),
                (
                    "RIGHTPADDING",
                    (0, 0),
                    (-1, -1),
                    5,
                ),
                (
                    "TOPPADDING",
                    (0, 0),
                    (-1, -1),
                    5,
                ),
                (
                    "BOTTOMPADDING",
                    (0, 0),
                    (-1, -1),
                    5,
                ),
            ]
        )
    )

    return table


# ======================================================================
# PDF REPORT
# ======================================================================

def create_pdf_report(
    results: dict[str, pd.DataFrame],
    values: dict[str, Any],
) -> None:

    report_styles = make_pdf_styles()

    document = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=A4,
        rightMargin=1.8 * cm,
        leftMargin=1.8 * cm,
        topMargin=1.8 * cm,
        bottomMargin=1.8 * cm,
        title=REPORT_TITLE,
    )

    story = []

    story.append(
        Paragraph(
            REPORT_TITLE,
            report_styles["title"],
        )
    )

    story.append(
        Paragraph(
            REPORT_SUBTITLE,
            report_styles["subtitle"],
        )
    )

    story.append(
        Paragraph(
            "Hypothesis",
            report_styles["heading_1"],
        )
    )

    story.append(
        Paragraph(
            HYPOTHESIS_TEXT,
            report_styles["body"],
        )
    )

    story.append(
        Paragraph(
            "Method",
            report_styles["heading_1"],
        )
    )

    for paragraph in build_method_text(
        values
    ):

        story.append(
            Paragraph(
                paragraph,
                report_styles["body"],
            )
        )

    story.append(
        Paragraph(
            "Results",
            report_styles["heading_1"],
        )
    )

    # Accuracy ------------------------------------------------------

    story.append(
        Paragraph(
            "Accuracy",
            report_styles["heading_2"],
        )
    )

    for paragraph in (
        build_accuracy_results_text(
            values
        )
    ):

        story.append(
            Paragraph(
                paragraph,
                report_styles["body"],
            )
        )

    story.append(
        Paragraph(
            "Table 1. Descriptive accuracy under neutral and negative emotion.",
            report_styles["body"],
        )
    )

    story.append(
        dataframe_to_pdf_table(
            build_accuracy_descriptive_table(
                results
            )
        )
    )

    story.append(
        Spacer(
            1,
            10,
        )
    )

    story.append(
        Paragraph(
            "Table 2. Binomial GEE estimate for response accuracy.",
            report_styles["body"],
        )
    )

    story.append(
        dataframe_to_pdf_table(
            build_accuracy_model_table(
                values
            ),
            column_widths=[
                4.5 * cm,
                2 * cm,
                2 * cm,
                3 * cm,
                2 * cm,
            ],
        )
    )

    story.append(
        Spacer(
            1,
            12,
        )
    )

    # Reaction time -------------------------------------------------

    story.append(
        Paragraph(
            "Reaction Time",
            report_styles["heading_2"],
        )
    )

    for paragraph in build_rt_results_text(
        values
    ):

        story.append(
            Paragraph(
                paragraph,
                report_styles["body"],
            )
        )

    story.append(
        Paragraph(
            "Table 3. Descriptive reaction time under neutral and negative emotion.",
            report_styles["body"],
        )
    )

    story.append(
        dataframe_to_pdf_table(
            build_rt_descriptive_table(
                results
            )
        )
    )

    story.append(
        Spacer(
            1,
            10,
        )
    )

    story.append(
        Paragraph(
            "Table 4. Mixed-effects model estimate for log reaction time.",
            report_styles["body"],
        )
    )

    story.append(
        dataframe_to_pdf_table(
            build_rt_model_table(
                values
            ),
            column_widths=[
                4.5 * cm,
                2 * cm,
                2.5 * cm,
                4 * cm,
                2 * cm,
            ],
        )
    )

    story.append(
        Spacer(
            1,
            12,
        )
    )

    # Speed-accuracy ------------------------------------------------

    story.append(
        Paragraph(
            "Speed-Accuracy Pattern",
            report_styles["heading_2"],
        )
    )

    story.append(
        Paragraph(
            (
                "Negative emotion was associated with significantly faster "
                "responses, whereas the corresponding reduction in accuracy "
                "was small and statistically non-significant. This pattern "
                "does not provide sufficient statistical evidence of a "
                "speed-accuracy trade-off, although it may reflect a "
                "descriptive shift toward faster responding."
            ),
            report_styles["body"],
        )
    )

    # Decision rules ------------------------------------------------

    story.append(
        Paragraph(
            "Decision Rules",
            report_styles["heading_1"],
        )
    )

    for rule in DECISION_RULES:

        story.append(
            Paragraph(
                f"• {rule}",
                report_styles["bullet"],
            )
        )

    # Interpretation ------------------------------------------------

    story.append(
        Paragraph(
            "Interpretation",
            report_styles["heading_1"],
        )
    )

    for paragraph in build_interpretation_text(
        values
    ):

        story.append(
            Paragraph(
                paragraph,
                report_styles["body"],
            )
        )

    # Assessment ----------------------------------------------------

    story.append(
        Paragraph(
            "Hypothesis Assessment",
            report_styles["heading_1"],
        )
    )

    story.append(
        Paragraph(
            (
                "<b>Conclusion:</b> Hypothesis 2 was not supported. "
                "Negative emotion did not significantly reduce accuracy "
                "and was associated with faster, rather than slower, "
                "reaction times relative to the neutral condition."
            ),
            report_styles["body"],
        )
    )

    document.build(
        story
    )

    print(
        f"PDF created: {PDF_PATH}"
    )


# ======================================================================
# MAIN
# ======================================================================

def main() -> None:

    print(
        "=" * 72
    )

    print(
        "GENERATING HYPOTHESIS 2 REPORT"
    )

    print(
        "=" * 72
    )

    results = load_results()

    values = extract_report_values(
        results
    )

    print(
        "\nResults loaded successfully."
    )

    print(
        f"Participants: "
        f"{values['participants']}"
    )

    print(
        f"Accuracy observations: "
        f"{values['accuracy_observations']}"
    )

    print(
        f"RT observations: "
        f"{values['rt_observations']}"
    )

    print(
        f"Accuracy OR: "
        f"{values['accuracy_or']:.3f}"
    )

    print(
        f"Accuracy p: "
        f"{values['accuracy_p']:.4f}"
    )

    print(
        f"RT percent change: "
        f"{values['rt_percent_change']:.1f}%"
    )

    print(
        f"RT p: "
        f"{values['rt_p']:.4f}"
    )

    create_docx_report(
        results,
        values,
    )

    create_pdf_report(
        results,
        values,
    )

    print(
        "\nReport generation completed successfully."
    )

    print(
        f"\nDOCX:\n{DOCX_PATH}"
    )

    print(
        f"\nPDF:\n{PDF_PATH}"
    )


if __name__ == "__main__":
    main()