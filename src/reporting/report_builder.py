"""
Generic DOCX and PDF report builder.

This module contains no hypothesis-specific statistical logic.
"""

from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path
from typing import Sequence

import pandas as pd

from docx import Document
from docx.enum.text import (
    WD_ALIGN_PARAGRAPH,
)
from docx.shared import Inches

from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.platypus import (
    Paragraph,
    SimpleDocTemplate,
    Spacer,
)

from .report_styles import (
    configure_docx_styles,
    create_pdf_styles,
)

from .report_tables import (
    add_docx_table,
    dataframe_to_pdf_table,
)


# ======================================================================
# REPORT SPECIFICATIONS
# ======================================================================

@dataclass
class ReportTable:
    """
    Table specification used by both report formats.
    """

    caption: str

    data: pd.DataFrame

    pdf_column_widths: (
        Sequence | None
    ) = None


@dataclass
class ReportSubsection:
    """
    Report subsection.
    """

    title: str

    paragraphs: list[str] = field(
        default_factory=list
    )

    tables: list[ReportTable] = field(
        default_factory=list
    )


@dataclass
class ReportSection:
    """
    Main report section.
    """

    title: str

    paragraphs: list[str] = field(
        default_factory=list
    )

    subsections: list[
        ReportSubsection
    ] = field(
        default_factory=list
    )

    tables: list[ReportTable] = field(
        default_factory=list
    )

    bullets: list[str] = field(
        default_factory=list
    )


@dataclass
class ReportContent:
    """
    Complete hypothesis-report specification.
    """

    title: str

    subtitle: str

    hypothesis: str

    sections: list[
        ReportSection
    ] = field(
        default_factory=list
    )


# ======================================================================
# BUILDER
# ======================================================================

class ScientificReportBuilder:
    """
    Build standardized scientific reports in DOCX and PDF formats.
    """

    def __init__(
        self,
        content: ReportContent,
    ) -> None:

        self.content = content

    # ------------------------------------------------------------------
    # DOCX
    # ------------------------------------------------------------------

    def build_docx(
        self,
        output_path: str | Path,
    ) -> Path:

        output_path = Path(
            output_path
        )

        output_path.parent.mkdir(
            parents=True,
            exist_ok=True,
        )

        document = Document()

        configure_docx_styles(
            document
        )

        section = (
            document.sections[0]
        )

        section.top_margin = Inches(
            0.8
        )

        section.bottom_margin = Inches(
            0.8
        )

        section.left_margin = Inches(
            0.9
        )

        section.right_margin = Inches(
            0.9
        )

        # Title
        title = document.add_heading(
            self.content.title,
            level=0,
        )

        title.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

        subtitle = (
            document.add_paragraph(
                self.content.subtitle
            )
        )

        subtitle.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

        # Hypothesis
        document.add_heading(
            "Hypothesis",
            level=1,
        )

        document.add_paragraph(
            self.content.hypothesis
        )

        # Sections
        for report_section in (
            self.content.sections
        ):

            document.add_heading(
                report_section.title,
                level=1,
            )

            for paragraph in (
                report_section.paragraphs
            ):

                document.add_paragraph(
                    paragraph
                )

            for bullet in (
                report_section.bullets
            ):

                document.add_paragraph(
                    bullet,
                    style="List Bullet",
                )

            for table in (
                report_section.tables
            ):

                document.add_paragraph(
                    table.caption
                )

                add_docx_table(
                    document,
                    table.data,
                )

                document.add_paragraph()

            # Subsections
            for subsection in (
                report_section.subsections
            ):

                document.add_heading(
                    subsection.title,
                    level=2,
                )

                for paragraph in (
                    subsection.paragraphs
                ):

                    document.add_paragraph(
                        paragraph
                    )

                for table in (
                    subsection.tables
                ):

                    document.add_paragraph(
                        table.caption
                    )

                    add_docx_table(
                        document,
                        table.data,
                    )

                    document.add_paragraph()

        document.save(
            output_path
        )

        return output_path

    # ------------------------------------------------------------------
    # PDF
    # ------------------------------------------------------------------

    def build_pdf(
        self,
        output_path: str | Path,
    ) -> Path:

        output_path = Path(
            output_path
        )

        output_path.parent.mkdir(
            parents=True,
            exist_ok=True,
        )

        styles = (
            create_pdf_styles()
        )

        document = (
            SimpleDocTemplate(
                str(output_path),
                pagesize=A4,
                rightMargin=1.8 * cm,
                leftMargin=1.8 * cm,
                topMargin=1.8 * cm,
                bottomMargin=1.8 * cm,
                title=self.content.title,
            )
        )

        story = []

        story.append(
            Paragraph(
                self.content.title,
                styles["title"],
            )
        )

        story.append(
            Paragraph(
                self.content.subtitle,
                styles["subtitle"],
            )
        )

        # Hypothesis
        story.append(
            Paragraph(
                "Hypothesis",
                styles["heading1"],
            )
        )

        story.append(
            Paragraph(
                self.content.hypothesis,
                styles["body"],
            )
        )

        # Sections
        for report_section in (
            self.content.sections
        ):

            story.append(
                Paragraph(
                    report_section.title,
                    styles["heading1"],
                )
            )

            for paragraph in (
                report_section.paragraphs
            ):

                story.append(
                    Paragraph(
                        paragraph,
                        styles["body"],
                    )
                )

            for bullet in (
                report_section.bullets
            ):

                story.append(
                    Paragraph(
                        f"• {bullet}",
                        styles["bullet"],
                    )
                )

            for table in (
                report_section.tables
            ):

                story.append(
                    Paragraph(
                        table.caption,
                        styles["caption"],
                    )
                )

                story.append(
                    dataframe_to_pdf_table(
                        table.data,
                        column_widths=(
                            table.pdf_column_widths
                        ),
                    )
                )

                story.append(
                    Spacer(
                        1,
                        10,
                    )
                )

            # Subsections
            for subsection in (
                report_section.subsections
            ):

                story.append(
                    Paragraph(
                        subsection.title,
                        styles["heading2"],
                    )
                )

                for paragraph in (
                    subsection.paragraphs
                ):

                    story.append(
                        Paragraph(
                            paragraph,
                            styles["body"],
                        )
                    )

                for table in (
                    subsection.tables
                ):

                    story.append(
                        Paragraph(
                            table.caption,
                            styles["caption"],
                        )
                    )

                    story.append(
                        dataframe_to_pdf_table(
                            table.data,
                            column_widths=(
                                table.pdf_column_widths
                            ),
                        )
                    )

                    story.append(
                        Spacer(
                            1,
                            10,
                        )
                    )

        document.build(
            story
        )

        return output_path